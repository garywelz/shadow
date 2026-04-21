from __future__ import annotations

from io import BytesIO
from typing import Any


def export_docx(manuscript: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title = manuscript.get("title") or "Manuscript"
    doc.add_heading(title, level=0)

    chapters = manuscript.get("chapters") or []
    for idx, ch in enumerate(chapters, start=1):
        ch_title = ch.get("title") or f"Chapter {idx}"
        doc.add_heading(ch_title, level=1)
        for seg in ch.get("segments") or []:
            text = (seg.get("text") or "").strip()
            if not text:
                continue
            for para in text.split("\n\n"):
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

